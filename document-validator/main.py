"""
Сервис проверки документов (DOCX) по шаблону.
Запуск: uvicorn main:app --reload --port 8082
"""
import re
from typing import Optional, List

from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

app = FastAPI(title="Document Validator", version="1.0.0")


class TemplateCriteria(BaseModel):
    """Критерии шаблона для проверки."""
    fileFormat: str = ".doc или .docx"
    font: str = "Times New Roman"
    fontSize: str = "14 пт"
    lineSpacing: str = "1.5"
    minPages: str = "10"
    minSources: str = "7"
    hasTitlePage: bool = True
    hasToc: bool = True
    hasIntroduction: bool = True
    hasMainPart: bool = True
    hasConclusion: bool = True
    hasBibliography: bool = True
    hasAppendices: bool = False
    illNumbering: Optional[str] = None
    figurePosition: Optional[str] = None
    figureCaption: Optional[str] = None
    tableTitle: Optional[str] = None
    tablePosition: Optional[str] = None


class ValidationError(BaseModel):
    code: str
    message: str
    passed: bool = False
    expected: Optional[str] = None
    actual: Optional[str] = None
    level: Optional[str] = None
    title: Optional[str] = None


class ValidationResult(BaseModel):
    valid: bool
    errors: List[ValidationError] = []
    warnings: List[ValidationError] = []
    details: dict = {}


def parse_int(s: str, default: int = 0) -> int:
    """Извлекает число из строки (например '10+ стр.' -> 10, '14 пт' -> 14)."""
    if not s:
        return default
    match = re.search(r"\d+", str(s))
    return int(match.group()) if match else default


def check_docx(file_content: bytes, criteria: TemplateCriteria) -> ValidationResult:
    """Проверяет DOCX-документ по критериям шаблона."""
    errors: List[ValidationError] = []
    warnings: List[ValidationError] = []
    details: dict = {}

    try:
        from io import BytesIO
        doc = Document(BytesIO(file_content))
    except Exception as e:
        return ValidationResult(
            valid=False,
            errors=[ValidationError(code="PARSE_ERROR", message=f"Не удалось открыть документ: {e}")],
            details={},
        )

    # --- Формат файла ---
    details["format"] = "DOCX"
    # (проверка расширения делается на клиенте)

    # --- Шрифт и размер ---
    expected_font = (criteria.font or "Times New Roman").split(",")[0].strip()
    expected_pt = parse_int(criteria.fontSize, 14)
    if expected_pt <= 0:
        expected_pt = 14

    font_ok = False
    size_ok = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                f = run.font.name.split(",")[0].strip()
                if expected_font.lower() in f.lower():
                    font_ok = True
            if run.font.size:
                pt = run.font.size.pt if run.font.size else 0
                if abs(pt - expected_pt) <= 1:
                    size_ok = True
        if font_ok and size_ok:
            break

    if not font_ok:
        errors.append(ValidationError(
            code="FONT",
            message=f"Требуется шрифт: {expected_font}",
            passed=False,
        ))
    else:
        details["font"] = "OK"

    if not size_ok:
        errors.append(ValidationError(
            code="FONT_SIZE",
            message=f"Требуется размер шрифта: {expected_pt} пт",
            passed=False,
        ))
    else:
        details["fontSize"] = "OK"

    # --- Межстрочный интервал ---
    expected_spacing = parse_int(criteria.lineSpacing.replace(",", ".").replace("1.5", "1"))
    if "1.5" in str(criteria.lineSpacing):
        expected_spacing = 15  # 1.5 строки
    spacing_ok = False
    for para in doc.paragraphs:
        if para.paragraph_format.line_spacing is not None:
            spacing_ok = True
            break
    if not spacing_ok:
        warnings.append(ValidationError(
            code="LINE_SPACING",
            message=f"Проверьте межстрочный интервал: {criteria.lineSpacing}",
            passed=False,
        ))
    else:
        details["lineSpacing"] = "OK"

    # --- Минимальное количество страниц (приблизительно) ---
    min_pages = parse_int(criteria.minPages, 10)
    # ~3000 символов на страницу A4, 14pt, 1.5 интервал
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    est_pages = max(1, total_chars // 2800)
    details["estimatedPages"] = est_pages
    details["minPages"] = min_pages

    if est_pages < min_pages:
        errors.append(ValidationError(
            code="MIN_PAGES",
            message=f"Объём недостаточен: ~{est_pages} стр. (требуется минимум {min_pages})",
            passed=False,
        ))

    # --- Минимальное количество источников ---
    min_sources = parse_int(criteria.minSources, 7)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Ищем "список литературы", "библиограф", номера [1], [2] и т.д.
    source_patterns = [
        r"\[\s*\d+\s*\]",
        r"\d+\.\s+[А-Яа-яЁёA-Za-z].*\.\s*(—|М\.|СПб\.|Л\.|М,|СПб,)",
        r"(?i)(литератур|источник|библиограф)",
    ]
    sources_found = set()
    for pat in source_patterns:
        for m in re.finditer(pat, full_text):
            sources_found.add(m.group()[:50])
    # Также считаем строки в типичном блоке списка литературы
    lit_count = len(re.findall(r"\[\d+\]", full_text))
    if lit_count > 0:
        sources_count = lit_count
    else:
        sources_count = max(len(sources_found), full_text.lower().count("— м.") + 1)
    details["sourcesFound"] = sources_count
    details["minSources"] = min_sources

    if sources_count < min_sources and criteria.hasBibliography:
        warnings.append(ValidationError(
            code="MIN_SOURCES",
            message=f"Источников в списке литературы: ~{sources_count} (рекомендуется {min_sources}+)",
            passed=False,
        ))

    # --- Структура документа ---
    text_lower = full_text.lower()
    structure_checks = [
        ("hasTitlePage", criteria.hasTitlePage, ["титульн", "title", "министерство", "федеральное"]),
        ("hasToc", criteria.hasToc, ["оглавление", "содержание", "contents"]),
        ("hasIntroduction", criteria.hasIntroduction, ["введение"]),
        ("hasMainPart", criteria.hasMainPart, ["глава", "раздел", "1.", "1 ", "основная часть"]),
        ("hasConclusion", criteria.hasConclusion, ["заключение"]),
        ("hasBibliography", criteria.hasBibliography, ["литератур", "источник", "библиограф", "список использован"]),
        ("hasAppendices", criteria.hasAppendices, ["приложение", "приложения"]),
    ]
    for key, required, keywords in structure_checks:
        found = any(kw in text_lower for kw in keywords)
        details[key] = "OK" if found else "NOT_FOUND"
        if required and not found:
            name = {
                "hasTitlePage": "Титульный лист",
                "hasToc": "Оглавление",
                "hasIntroduction": "Введение",
                "hasMainPart": "Основная часть",
                "hasConclusion": "Заключение",
                "hasBibliography": "Список литературы",
                "hasAppendices": "Приложения",
            }.get(key, key)
            errors.append(ValidationError(
                code=key.upper(),
                message=f"Отсутствует раздел: {name}",
                passed=False,
            ))

    # --- Иллюстрации и таблицы (по триггерным словам "Рис.", "Рисунок", "Таблица"), уровень warning ---
    fig_triggers = ("рис.", "рисунок")
    table_trigger = "таблица"
    paragraphs = doc.paragraphs

    if criteria.figurePosition:
        expected_pos = (criteria.figurePosition or "").strip().lower()
        center_ok = "центр" in expected_pos or "по центру" in expected_pos
        for i, para in enumerate(paragraphs):
            t = para.text.strip()
            if not t:
                continue
            tl = t.lower()
            if any(tr in tl for tr in fig_triggers) and ("рис" in tl or "рисунок" in tl):
                alignment = para.paragraph_format.alignment
                is_center = alignment == WD_ALIGN_PARAGRAPH.CENTER if alignment is not None else False
                if center_ok and not is_center:
                    actual = "По центру" if is_center else (
                        "Слева" if alignment == WD_ALIGN_PARAGRAPH.LEFT else
                        "Справа" if alignment == WD_ALIGN_PARAGRAPH.RIGHT else "Не задано"
                    )
                    warnings.append(ValidationError(
                        code="FIGURE_POSITION",
                        message="Позиция рисунка не по центру",
                        passed=False,
                        expected=criteria.figurePosition,
                        actual=actual,
                        level="warning",
                        title="Позиция рисунка",
                    ))
                elif not center_ok and is_center and "слева" in expected_pos and alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    warnings.append(ValidationError(
                        code="FIGURE_POSITION",
                        message="Позиция рисунка не соответствует шаблону",
                        passed=False,
                        expected=criteria.figurePosition,
                        actual="По центру",
                        level="warning",
                        title="Позиция рисунка",
                    ))
                break

    if criteria.figureCaption:
        for i, para in enumerate(paragraphs):
            t = para.text.strip()
            if not t:
                continue
            if any(tr in t.lower() for tr in fig_triggers):
                expected_cap = (criteria.figureCaption or "").strip().lower()
                if "под" in expected_cap and i + 1 < len(paragraphs):
                    next_t = paragraphs[i + 1].text.strip().lower()
                    if next_t and ("рис" in next_t or "рисунок" in next_t):
                        details["figureCaption"] = "OK"
                    else:
                        details["figureCaption"] = "CHECK"
                break

    if criteria.tableTitle:
        full_text_lower = full_text.lower()
        if table_trigger in full_text_lower:
            if re.search(r"таблица\s*\d+", full_text_lower):
                details["tableTitle"] = "OK"
            else:
                warnings.append(ValidationError(
                    code="TABLE_TITLE",
                    message="Заголовок таблицы: ожидается «Таблица» + номер",
                    passed=False,
                    expected=criteria.tableTitle,
                    actual="Не найдено",
                    level="warning",
                    title="Заголовок таблицы",
                ))

    if criteria.tablePosition and doc.tables:
        expected_table_pos = (criteria.tablePosition or "").strip().lower()
        center_ok = "центр" in expected_table_pos or "по центру" in expected_table_pos
        for tbl in doc.tables:
            try:
                tbl_align = getattr(tbl, "alignment", None)
                if tbl_align is not None:
                    is_center = tbl_align == WD_TABLE_ALIGNMENT.CENTER
                    if center_ok and not is_center:
                        warnings.append(ValidationError(
                            code="TABLE_POSITION",
                            message="Таблица не по центру страницы",
                            passed=False,
                            expected=criteria.tablePosition,
                            actual="Слева" if tbl_align == WD_TABLE_ALIGNMENT.LEFT else "Другое",
                            level="warning",
                            title="Позиция таблицы",
                        ))
                    break
            except (AttributeError, TypeError):
                pass

    if criteria.illNumbering:
        fig_nums = re.findall(r"(?:рис\.|рисунок)\s*(\d+)", full_text_lower, re.IGNORECASE)
        if fig_nums:
            try:
                nums = [int(n) for n in fig_nums]
                sequential = nums == list(range(1, len(nums) + 1)) or nums == list(range(min(nums), min(nums) + len(nums)))
                if not sequential and "сквозная" in (criteria.illNumbering or "").lower():
                    warnings.append(ValidationError(
                        code="ILL_NUMBERING",
                        message="Нумерация иллюстраций не сквозная",
                        passed=False,
                        expected=criteria.illNumbering,
                        actual=", ".join(str(n) for n in nums[:5]) + ("..." if len(nums) > 5 else ""),
                        level="warning",
                        title="Нумерация иллюстраций",
                    ))
                else:
                    details["illNumbering"] = "OK"
            except (ValueError, TypeError):
                pass

    return ValidationResult(
        valid=len(errors) == 0,
        errors=errors,
        warnings=warnings,
        details=details,
    )


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/validate")
async def validate_document(
    file: UploadFile = File(...),
    template: Optional[str] = Form(None),
):
    """
    Проверяет загруженный документ по шаблону.
    Поддерживаются форматы: .docx
    template — JSON строка с критериями (TemplateCriteria). Если не указан — используется шаблон по умолчанию.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Файл не выбран")

    ext = (file.filename or "").lower().split(".")[-1]
    if ext not in ("docx",):
        raise HTTPException(
            status_code=400,
            detail="Поддерживается только формат DOCX. PDF и DOC пока не поддерживаются.",
        )

    content = await file.read()
    if len(content) < 100:
        raise HTTPException(status_code=400, detail="Файл слишком маленький или повреждён")

    import json
    criteria = TemplateCriteria()
    if template:
        try:
            data = json.loads(template)
            criteria = TemplateCriteria(**data)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Некорректный шаблон: {e}")

    result = check_docx(content, criteria)
    return result.model_dump()
