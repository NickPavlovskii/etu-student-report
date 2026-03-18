#!/usr/bin/env python3
"""Generate Word document for Student Works analysis report."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        heading_style.font.size = Pt(18)
    elif level == 2:
        heading_style.font.size = Pt(16)
    else:
        heading_style.font.size = Pt(14)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '4472C4',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
            if r_idx % 2 == 1:
                from docx.oxml.ns import qn
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'D6E4F0',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elm)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ==================== TITLE PAGE ====================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ИС "Студенческие работы"')
run.font.size = Pt(26)
run.bold = True
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Аналитический отчёт')
run.font.size = Pt(20)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Анализ кода, бизнес-модель, пользовательские сценарии и схемы')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.italic = True

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
doc.add_heading('Содержание', level=1)
toc_items = [
    '1. Анализ кода и архитектуры',
    '   1.1 Стек технологий',
    '   1.2 Архитектура приложения',
    '   1.3 Модульная структура фронтенда',
    '   1.4 Ключевые API-эндпоинты',
    '2. План действий пользователя (User Flows)',
    '   2.1 Роли в системе',
    '   2.2 Загрузка студенческой работы',
    '   2.3 Просмотр аналитики',
    '   2.4 Управление системой (администратор)',
    '   2.5 Работа с архивом',
    '   2.6 Настройка шаблонов валидации',
    '3. Бизнес-модель',
    '   3.1 Business Model Canvas',
    '   3.2 Ценностное предложение',
    '   3.3 Цепочка создания ценности',
    '4. Схемы и диаграммы (см. Draw.io файлы)',
    'Приложение: Ключевые метрики системы',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_page_break()

# ==================== SECTION 1 ====================
doc.add_heading('1. Анализ кода и архитектуры', level=1)

doc.add_heading('1.1 Стек технологий', level=2)
doc.add_paragraph('Система построена на современном стеке веб-технологий:')

add_table(
    ['Слой', 'Технологии'],
    [
        ['Frontend', 'Vue 3, TypeScript, Vite'],
        ['UI-библиотеки', 'Vuetify 3 (Material Design 3), Element Plus'],
        ['Стилизация', 'TailwindCSS, SASS'],
        ['HTTP-клиент', 'Axios'],
        ['Маршрутизация', 'Vue Router 4'],
        ['Экспорт', 'XLSX (Excel)'],
        ['Backend', 'Java Spring Boot (порт 8081)'],
        ['Валидация документов', 'Java-бэкенд + опциональный Python FastAPI'],
        ['Деплой', 'Docker, Nginx, Docker Compose'],
    ],
    col_widths=[5, 12]
)

doc.add_heading('1.2 Архитектура приложения', level=2)
doc.add_paragraph(
    'Приложение построено по архитектуре SPA (Single Page Application) с разделением '
    'на клиентскую и серверную части. Клиент реализован на Vue 3 с TypeScript, '
    'серверная часть — на Java Spring Boot. Взаимодействие осуществляется через REST API '
    '(JSON и FormData для загрузки файлов).'
)

p = doc.add_paragraph()
run = p.add_run('Клиентская часть ')
run.bold = True
p.add_run(
    'состоит из 7 функциональных модулей: авторизация, дисциплины, детали дисциплины '
    '(загрузка работ), архив, аналитика, настройки и администрирование. '
    'Общие компоненты (кнопки, карточки, боковая панель) вынесены в отдельный слой.'
)

p = doc.add_paragraph()
run = p.add_run('Серверная часть ')
run.bold = True
p.add_run(
    'предоставляет API для аутентификации, управления отчётами, валидации документов, '
    'управления шаблонами, административных операций и аналитики.'
)

doc.add_paragraph(
    'Подробная диаграмма компонентов доступна в файле Draw.io: diagrams/component-diagram.drawio'
)

doc.add_heading('1.3 Модульная структура фронтенда', level=2)

add_table(
    ['Директория', 'Назначение'],
    [
        ['src/api/', 'REST-клиенты для каждого домена (auth, disciplines, archive и др.)'],
        ['src/modules/auth/', 'Страницы авторизации (по фамилии, по ETU ID)'],
        ['src/modules/disciplines/', 'Список дисциплин преподавателя (дашборд)'],
        ['src/modules/discipline/', 'Детали дисциплины, загрузка работ, валидация'],
        ['src/modules/archive/', 'Архив работ с фильтрацией и экспортом'],
        ['src/modules/analytics/', 'Аналитика и статистика'],
        ['src/modules/settings/', 'Настройки шаблонов валидации'],
        ['src/modules/admin/', 'Панель администратора (персонал, аудит)'],
        ['src/components/', 'Переиспользуемые UI-компоненты'],
        ['src/composables/', 'Композиционные функции (бизнес-логика)'],
        ['src/types/', 'TypeScript-типы и интерфейсы'],
        ['src/utils/', 'Утилиты (форматирование дат, инициалы и др.)'],
    ],
    col_widths=[5, 12]
)

doc.add_heading('1.4 Ключевые API-эндпоинты', level=2)

add_table(
    ['Группа', 'Метод', 'Путь', 'Описание'],
    [
        ['Auth', 'GET', '/api/auth/login?lastName=', 'Авторизация по фамилии'],
        ['Дисциплины', 'GET', '/api/teachers/{lastName}/disciplines/cards', 'Карточки дисциплин'],
        ['Детали', 'GET', '/api/teachers/{lastName}/discipline/{id}', 'Информация о дисциплине'],
        ['Студенты', 'GET', '/api/teachers/{lastName}/discipline/{id}/students', 'Список студентов'],
        ['Отчёты', 'GET', '/api/teachers/{lastName}/discipline/{id}/reports', 'Загруженные работы'],
        ['Загрузка', 'POST', '/api/teachers/{lastName}/discipline/{id}/reports', 'Загрузить работу'],
        ['Скачать', 'GET', '/api/reports/{id}/download', 'Скачать отчёт'],
        ['Валидация', 'POST', '/api/validate', 'Проверить документ'],
        ['Шаблоны', 'CRUD', '/api/templates', 'Управление шаблонами'],
        ['Архив', 'GET', '/api/archive/reports', 'Архив всех работ'],
        ['Аналитика', 'GET', '/api/teachers/{lastName}/stats/*', 'Статистика преподавателя'],
        ['Админ', 'GET/PUT', '/api/admin/*', 'Административные операции'],
        ['Аудит', 'GET/POST', '/api/admin/audit-log', 'Журнал аудита и откат'],
    ],
    col_widths=[2.5, 1.5, 7, 5]
)

doc.add_page_break()

# ==================== SECTION 2 ====================
doc.add_heading('2. План действий пользователя (User Flows)', level=1)

doc.add_heading('2.1 Роли в системе', level=2)

doc.add_paragraph('В системе предусмотрены две роли пользователей:')

add_table(
    ['Возможность', 'Преподаватель', 'Администратор'],
    [
        ['Просмотр своих дисциплин', '✓', '✓'],
        ['Просмотр всех дисциплин кафедры', '✗', '✓'],
        ['Загрузка студенческих работ', '✓', '✓'],
        ['Скачивание работ', '✓', '✓'],
        ['Валидация документов', '✓', '✓'],
        ['Просмотр архива', '✓', '✓'],
        ['Экспорт в Excel', '✓', '✓'],
        ['Личная аналитика', '✓', '✓'],
        ['Аналитика по кафедре', '✗', '✓'],
        ['Настройка шаблонов валидации', '✓', '✓'],
        ['Управление персоналом', '✗', '✓'],
        ['Управление ролями', '✗', '✓'],
        ['Журнал аудита', '✗', '✓'],
        ['Откат действий', '✗', '✓'],
    ],
    col_widths=[7, 3, 3]
)

doc.add_heading('2.2 Сценарий: Загрузка студенческой работы', level=2)

steps = [
    ('Шаг 1', 'Авторизация', 'Преподаватель входит в систему по фамилии или ETU ID'),
    ('Шаг 2', 'Выбор дисциплины', 'На дашборде отображаются все дисциплины. Преподаватель выбирает нужную'),
    ('Шаг 3', 'Страница дисциплины', 'Отображается список групп и студентов с прогрессом загрузки'),
    ('Шаг 4', 'Открытие формы загрузки', 'Нажатие кнопки "Загрузить работу" открывает модальное окно'),
    ('Шаг 5', 'Заполнение метаданных', 'Выбор группы, студента, типа контроля, типа работы, темы'),
    ('Шаг 6', 'Прикрепление файла', 'Загрузка файла (PDF, DOCX, DOC) через drag-and-drop или выбор'),
    ('Шаг 7', 'Автоматическая валидация', 'При включённой опции система проверяет документ по шаблону'),
    ('Шаг 8', 'Результат', 'Отображается % соответствия, список ошибок и предупреждений'),
]

add_table(
    ['Этап', 'Действие', 'Описание'],
    steps,
    col_widths=[2, 4, 10]
)

doc.add_heading('2.3 Сценарий: Просмотр аналитики', level=2)

steps = [
    ('Шаг 1', 'Переход в раздел "Аналитика" через боковое меню'),
    ('Шаг 2', 'Выбор периода: неделя, месяц или учебный год'),
    ('Шаг 3', 'Просмотр KPI-карточек: всего работ, проверено, средний %, проблемных'),
    ('Шаг 4', 'Анализ графиков: динамика загрузок, соответствие по бакетам, частые ошибки'),
    ('Шаг 5', 'Просмотр таблиц: статистика по группам и дисциплинам'),
]

add_table(
    ['Этап', 'Действие'],
    steps,
    col_widths=[2, 14]
)

doc.add_heading('2.4 Сценарий: Управление системой (Администратор)', level=2)

p = doc.add_paragraph()
run = p.add_run('Панель администратора ')
run.bold = True
p.add_run('содержит три вкладки:')

items = [
    'Персонал — список преподавателей кафедры, поиск, изменение ролей (TEACHER/ADMIN)',
    'Дисциплины — обзор всех дисциплин кафедры',
    'Журнал аудита — история всех действий в системе с фильтрацией по типу действия, '
    'актору и дате. Возможность отката действий (загрузок, удалений)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.5 Сценарий: Работа с архивом', level=2)

steps = [
    ('Шаг 1', 'Переход в раздел "Архив" через боковое меню'),
    ('Шаг 2', 'Настройка фильтров: дисциплина, группа, преподаватель, тип работы, период'),
    ('Шаг 3', 'Поиск по ключевым словам'),
    ('Шаг 4', 'Просмотр таблицы результатов (студент, дисциплина, дата, статус, версия)'),
    ('Шаг 5', 'Скачивание отдельных работ (PDF/DOCX) или экспорт всего в Excel'),
]

add_table(
    ['Этап', 'Действие'],
    steps,
    col_widths=[2, 14]
)

doc.add_heading('2.6 Сценарий: Настройка шаблонов валидации', level=2)

doc.add_paragraph('Раздел "Настройки" позволяет создавать и редактировать шаблоны для автоматической проверки документов:')

criteria = [
    ('Формат файла', '.pdf, .docx, .doc'),
    ('Шрифт', 'Например, Times New Roman'),
    ('Размер шрифта', 'Например, 14 пт'),
    ('Межстрочный интервал', 'Например, 1.5'),
    ('Минимум страниц', 'Числовое значение'),
    ('Минимум источников', 'Числовое значение'),
    ('Титульный лист', 'Обязательный / необязательный'),
    ('Содержание', 'Обязательный / необязательный'),
    ('Введение', 'Обязательный / необязательный'),
    ('Основная часть', 'Обязательный / необязательный'),
    ('Заключение', 'Обязательный / необязательный'),
    ('Библиография', 'Обязательный / необязательный'),
    ('Приложения', 'Обязательный / необязательный'),
]

add_table(
    ['Критерий', 'Значение'],
    criteria,
    col_widths=[5, 11]
)

doc.add_page_break()

# ==================== SECTION 3 ====================
doc.add_heading('3. Бизнес-модель', level=1)

doc.add_heading('3.1 Business Model Canvas', level=2)

canvas = [
    ('Ключевые партнёры', 'ВУЗы, кафедры, ИТ-отдел университета'),
    ('Ценностное предложение', 'Единая система для сбора, валидации, архивации и аналитики студенческих работ'),
    ('Целевые сегменты', 'Преподаватели вузов, администрация кафедр, деканаты, учебные управления'),
    ('Каналы', 'Веб-интерфейс (SPA), интеграция с ЕТУ, API для внешних систем'),
    ('Ключевые ресурсы', 'Серверная инфраструктура, хранилище документов, сервис валидации'),
    ('Ключевые активности', 'Разработка платформы, обновление шаблонов под ГОСТ, техподдержка, интеграции'),
    ('Структура издержек', 'Хостинг/серверы, разработка и поддержка, хранение документов'),
    ('Потоки доходов', 'Внутренний продукт ВУЗа (бюджет), лицензирование для других ВУЗов'),
]

add_table(
    ['Блок Canvas', 'Описание'],
    canvas,
    col_widths=[5, 11]
)

doc.add_heading('3.2 Ценностное предложение (Value Proposition)', level=2)

vp = [
    ('Ручной сбор студенческих работ', 'Централизованная загрузка через веб', 'Экономия времени преподавателя'),
    ('Проверка оформления вручную', 'Автоматическая валидация по шаблону', 'Объективность и скорость проверки'),
    ('Отсутствие единого хранилища', 'Архив с поиском и фильтрацией', 'Доступ к истории работ'),
    ('Нет прозрачности процесса', 'Аналитика и статистика', 'Контроль качества учебного процесса'),
    ('Сложность аудита', 'Журнал действий с откатом', 'Подотчётность и безопасность'),
]

add_table(
    ['Проблема', 'Решение', 'Ценность'],
    vp,
    col_widths=[5, 6, 5]
)

doc.add_heading('3.3 Цепочка создания ценности', level=2)

chain = [
    ('1. Настройка шаблонов', 'Стандартизация требований к оформлению'),
    ('2. Загрузка работ', 'Удобство для преподавателей'),
    ('3. Автоматическая валидация', 'Объективная оценка оформления'),
    ('4. Хранение в архиве', 'Долгосрочное хранение и доступ'),
    ('5. Аналитика и отчёты', 'Принятие решений на основе данных'),
]

add_table(
    ['Этап', 'Создаваемая ценность'],
    chain,
    col_widths=[5, 11]
)

doc.add_page_break()

# ==================== SECTION 4 ====================
doc.add_heading('4. Схемы и диаграммы', level=1)

doc.add_paragraph(
    'Все схемы и диаграммы системы представлены в формате Draw.io и доступны '
    'в директории docs/diagrams/. Файлы можно открыть на сайте https://app.diagrams.net/ '
    'или в десктопном приложении Draw.io.'
)

diagrams_list = [
    ('component-diagram.drawio', 'Диаграмма компонентов системы', 'Общая архитектура: Frontend, Backend, инфраструктура'),
    ('state-diagram.drawio', 'Диаграмма состояний документа', 'Жизненный цикл студенческой работы: загрузка → валидация → архив'),
    ('er-diagram.drawio', 'ER-диаграмма (модель данных)', 'Сущности: Teacher, Discipline, Student, Report, Template, AuditLog'),
    ('sequence-diagram.drawio', 'Диаграмма последовательности', 'Процесс загрузки и валидации работы'),
    ('navigation-diagram.drawio', 'Карта навигации (Sitemap)', 'Структура страниц приложения'),
    ('deployment-diagram.drawio', 'Диаграмма развёртывания', 'Docker, Nginx, Spring Boot'),
    ('usecase-diagram.drawio', 'Диаграмма вариантов использования', 'Use Cases для преподавателя и администратора'),
    ('dfd-diagram.drawio', 'Диаграмма потоков данных (DFD)', 'Потоки данных между компонентами системы'),
]

add_table(
    ['Файл', 'Название', 'Описание'],
    diagrams_list,
    col_widths=[5, 5, 6]
)

doc.add_page_break()

# ==================== APPENDIX ====================
doc.add_heading('Приложение: Ключевые метрики системы', level=1)

metrics = [
    ('Всего загруженных работ', 'Общее количество загруженных студенческих работ'),
    ('Проверено по шаблону', 'Работы, прошедшие автоматическую валидацию'),
    ('Средний % соответствия', 'Среднее соответствие стандартам оформления'),
    ('Проблемных работ', 'Работы с критическими ошибками оформления'),
    ('Динамика загрузок', 'Ежемесячная активность загрузки работ'),
    ('Частые ошибки', 'Наиболее распространённые ошибки оформления'),
    ('Проблемные группы', 'Группы с наибольшим количеством ошибок'),
    ('Распределение соответствия', 'Гистограмма по бакетам % соответствия'),
]

add_table(
    ['Метрика', 'Описание'],
    metrics,
    col_widths=[5, 11]
)

# --- Save ---
output_path = os.path.join(os.path.dirname(__file__), 'analysis-report.docx')
doc.save(output_path)
print(f'Document saved: {output_path}')
